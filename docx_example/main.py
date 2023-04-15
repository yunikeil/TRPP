import datetime
from docx import Document


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


def generate_document(file_path, template_file, output_file):
    with open(template_file) as f:
        template_doc = Document(f)

    data = {
        "{{EMPLOYEE_POSITION}}": "программсит",
        "{{EMPLOYEE_FULLNAME}}": "Василий И.В.",
        "{{EMPLOYEE_SHORTNAME}}": "Вася",
        "{{START_DATE}}": "01.01.2022",
        "{{FINISH_DATE}}": "31.12.2022",
        "{{DURATION}}": "365",
        "{{DATE}}": datetime.datetime.now().strftime("%d.%m.%Y")
    }

    if file_path:
        with open(file_path) as f:
            lines = f.read().split("\n")
        if len(lines) > 0 and "{{" in lines[0]:
            for line in lines:
                if "{{" in line and "}}" in line:
                    key, value = line.split(":")
                    data[key.strip()] = value.strip()
        else:
            data = dict(zip(data.keys(), lines))

    for key, value in data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)
            for table in paragraph.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text(paragraph, key, value)

    template_doc.save(output_file)


if __name__ == "__main__":
    main()
